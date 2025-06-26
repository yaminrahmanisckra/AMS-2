from flask import Blueprint, render_template, redirect, url_for, flash, request, send_file
from flask_login import login_required, current_user
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Frame, PageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.lib.units import inch
from io import BytesIO
from .models import db, RSession, RStudent, RSubject, RMark, RCourseRegistration
from openpyxl import load_workbook
import io
from datetime import datetime
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, A4
import zipfile
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

result_management_bp = Blueprint('result_management', __name__, template_folder='templates/result_management')

def calculate_grade(total_marks, is_retake=False):
    grade_point = 0.0
    grade_letter = 'F'
    if total_marks >= 80: grade_point, grade_letter = 4.0, 'A+'
    elif total_marks >= 75: grade_point, grade_letter = 3.75, 'A'
    elif total_marks >= 70: grade_point, grade_letter = 3.5, 'A-'
    elif total_marks >= 65: grade_point, grade_letter = 3.25, 'B+'
    elif total_marks >= 60: grade_point, grade_letter = 3.0, 'B'
    elif total_marks >= 55: grade_point, grade_letter = 2.75, 'B-'
    elif total_marks >= 50: grade_point, grade_letter = 2.5, 'C+'
    elif total_marks >= 45: grade_point, grade_letter = 2.25, 'C'
    elif total_marks >= 40: grade_point, grade_letter = 2.0, 'D'
    if is_retake and grade_letter != 'F':
        if grade_letter == 'A+': grade_point, grade_letter = 3.75, 'A'
        elif grade_letter == 'A': grade_point, grade_letter = 3.5, 'A-'
        elif grade_letter == 'A-': grade_point, grade_letter = 3.25, 'B+'
        elif grade_letter == 'B+': grade_point, grade_letter = 3.0, 'B'
        elif grade_letter == 'B': grade_point, grade_letter = 2.75, 'B-'
        elif grade_letter == 'B-': grade_point, grade_letter = 2.5, 'C+'
        elif grade_letter == 'C+': grade_point, grade_letter = 2.25, 'C'
        elif grade_letter == 'C': grade_point, grade_letter = 2.0, 'D'
    return grade_point, grade_letter

def convert_to_roman(num):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

# --- Page Numbering Function ---
def _footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 9)
    page_number_text = f"Page {doc.page} of {doc.doc.page_count}"
    canvas.drawRightString(letter[0] - 40, 30, page_number_text)
    canvas.restoreState()

@result_management_bp.route('/')
@login_required
def index():
    sessions = RSession.query.filter_by(is_archived=False).order_by(RSession.created_at.desc()).all()
    return render_template('rm_index.html', sessions=sessions)

@result_management_bp.route('/archived')
@login_required
def archived_sessions():
    sessions = RSession.query.filter_by(is_archived=True).order_by(RSession.created_at.desc()).all()
    return render_template('rm_archive.html', sessions=sessions)

@result_management_bp.route('/archive_session/<int:session_id>', methods=['POST'])
@login_required
def archive_session(session_id):
    session = RSession.query.get_or_404(session_id)
    session.is_archived = True
    db.session.commit()
    flash(f'Session "{session.name}" has been archived.', 'success')
    return redirect(url_for('result_management.index'))

@result_management_bp.route('/unarchive_session/<int:session_id>', methods=['POST'])
@login_required
def unarchive_session(session_id):
    session = RSession.query.get_or_404(session_id)
    session.is_archived = False
    db.session.commit()
    flash(f'Session "{session.name}" has been unarchived.', 'success')
    return redirect(url_for('result_management.archived_sessions'))

@result_management_bp.route('/add_session', methods=['GET', 'POST'])
@login_required
def add_session():
    if request.method == 'POST':
        name = request.form.get('name')
        term = request.form.get('term')
        year = request.form.get('year')
        
        if name and term:
            new_session = RSession(name=name, term=term, year=year)
            db.session.add(new_session)
            db.session.commit()
            flash('Session added successfully!', 'success')
            return redirect(url_for('result_management.index'))
        else:
            flash('Session Name and Term are required.', 'danger')
            
    return render_template('rm_add_session.html')

@result_management_bp.route('/add_student/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_student(session_id):
    if request.method == 'POST':
        if 'excel_file' in request.files and request.files['excel_file'].filename != '':
            file = request.files['excel_file']
            if file and file.filename.endswith('.xlsx'):
                try:
                    wb = load_workbook(file)
                    ws = wb.active
                    added_count = 0
                    skipped_count = 0

                    # --- Optimization Start ---
                    # 1. Read all student IDs from the Excel file
                    all_student_ids_from_excel = []
                    rows_to_process = []
                    for row in ws.iter_rows(min_row=2, values_only=True):
                        student_data = (list(row) + [None]*5)[:5]
                        student_id = student_data[0]
                        if student_id:
                            all_student_ids_from_excel.append(str(student_id))
                            rows_to_process.append(student_data)

                    # 2. Find which of these students already exist in the DB in a single query
                    existing_students = db.session.query(RStudent.student_id).filter(
                        RStudent.session_id == session_id,
                        RStudent.student_id.in_(all_student_ids_from_excel)
                    ).all()
                    existing_student_ids = {str(s_id[0]) for s_id in existing_students}

                    # 3. Iterate and add only the new students
                    students_to_add = []
                    for student_id, name, year, discipline, school in rows_to_process:
                        if str(student_id) not in existing_student_ids:
                            students_to_add.append(RStudent(
                                student_id=str(student_id), name=name, year=year,
                                discipline=discipline, school=school, session_id=session_id
                            ))
                            added_count += 1
                        else:
                            skipped_count += 1
                    
                    if students_to_add:
                        db.session.bulk_save_objects(students_to_add)
                    # --- Optimization End ---
                    
                    db.session.commit()
                    flash(f'Successfully added {added_count} new students. Skipped {skipped_count} existing students.', 'success')
                except Exception as e:
                    flash(f'Error processing Excel file: {e}', 'danger')
                return redirect(url_for('result_management.add_student', session_id=session_id))
            else:
                flash('Invalid file type. Please upload a .xlsx file.', 'danger')
        else:
            student_id = request.form.get('student_id')
            name = request.form.get('name')
            if student_id and name:
                exists = RStudent.query.filter_by(student_id=student_id, session_id=session_id).first()
                if exists:
                    flash('Student with this ID already exists in this session.', 'danger')
                else:
                    student = RStudent(
                        student_id=student_id, name=name, session_id=session_id
                    )
                    db.session.add(student)
                    db.session.commit()
                    flash('Student added successfully!', 'success')
            else:
                flash('Student ID and Name are required for single add.', 'warning')
        return redirect(url_for('result_management.add_student', session_id=session_id))
    
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()
    session = RSession.query.get_or_404(session_id)
    return render_template('rm_add_student.html', session=session, students=students)

@result_management_bp.route('/edit_student/<int:student_id>', methods=['GET', 'POST'])
@login_required
def edit_student(student_id):
    student = RStudent.query.get_or_404(student_id)
    if request.method == 'POST':
        student.student_id = request.form['student_id']
        student.name = request.form['name']
        student.year = request.form.get('year')
        student.discipline = request.form.get('discipline')
        student.school = request.form.get('school')
        db.session.commit()
        flash('Student updated successfully!', 'success')
        return redirect(url_for('result_management.add_student', session_id=student.session_id))
    return render_template('rm_edit_student.html', student=student)

@result_management_bp.route('/delete_student/<int:student_id>', methods=['POST'])
@login_required
def delete_student(student_id):
    student = RStudent.query.get_or_404(student_id)
    session_id = student.session_id
    db.session.delete(student)
    db.session.commit()
    flash('Student deleted successfully!', 'success')
    return redirect(url_for('result_management.add_student', session_id=session_id))

@result_management_bp.route('/add_subject/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_subject(session_id):
    if request.method == 'POST':
        code = request.form['code']
        name = request.form['name']
        credit = float(request.form['credit'])
        subject_type = request.form['subject_type']
        dissertation_type = request.form.get('dissertation_type') if subject_type == 'Dissertation' else None
        
        subject = RSubject(
            code=code, name=name, credit=credit, subject_type=subject_type,
            dissertation_type=dissertation_type, session_id=session_id
        )
        db.session.add(subject)
        db.session.commit()
        flash('Subject added successfully!', 'success')
        return redirect(url_for('result_management.add_subject', session_id=session_id))
    subjects = RSubject.query.filter_by(session_id=session_id).all()
    session = RSession.query.get_or_404(session_id)
    return render_template('rm_add_subject.html', session=session, subjects=subjects)

@result_management_bp.route('/delete_subject/<int:subject_id>', methods=['POST'])
@login_required
def delete_subject(subject_id):
    subject = RSubject.query.get_or_404(subject_id)
    session_id = subject.session_id
    db.session.delete(subject)
    db.session.commit()
    flash('Subject deleted successfully!', 'success')
    return redirect(url_for('result_management.add_subject', session_id=session_id))

@result_management_bp.route('/add_marks/<int:session_id>', methods=['GET', 'POST'])
@login_required
def add_marks(session_id):
    session = RSession.query.get_or_404(session_id)
    subjects = RSubject.query.filter_by(session_id=session_id).all()
    
    selected_subject_id = request.args.get('subject_id', type=int)
    selected_subject = RSubject.query.get(selected_subject_id) if selected_subject_id else None

    # Only show students who are registered for the selected subject
    if selected_subject:
        registered_student_ids = db.session.query(RCourseRegistration.student_id).filter_by(subject_id=selected_subject.id).all()
        registered_student_ids = [sid for (sid,) in registered_student_ids]
        students = RStudent.query.filter(RStudent.id.in_(registered_student_ids)).order_by(RStudent.student_id).all()
    else:
        students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()
    
    marks_data = {}
    registrations_data = {} # To store retake status
    if selected_subject:
        for student in students:
            mark = RMark.query.filter_by(student_id=student.id, subject_id=selected_subject.id).first()
            marks_data[student.id] = mark
            
            registration = RCourseRegistration.query.filter_by(student_id=student.id, subject_id=selected_subject.id).first()
            registrations_data[student.id] = registration

    if request.method == 'POST':
        subject_id = request.form.get('subject_id', type=int)
        if not subject_id:
            flash('Please select a subject.', 'danger')
            return redirect(url_for('result_management.add_marks', session_id=session_id))

        subject = RSubject.query.get_or_404(subject_id)

        for student in students:
            existing_mark = RMark.query.filter_by(student_id=student.id, subject_id=subject.id).first()
            if existing_mark is None:
                existing_mark = RMark(student_id=student.id, subject_id=subject.id)
                db.session.add(existing_mark)
            
            # Check if the student is registered for this course as a retake
            registration = RCourseRegistration.query.filter_by(student_id=student.id, subject_id=subject.id).first()
            is_retake = registration.is_retake if registration else False
            existing_mark.is_retake = is_retake

            total_marks = 0
            if subject.subject_type == 'Theory' or subject.subject_type == 'Theory (UG)':
                attendance = request.form.get(f'attendance_{student.id}')
                continuous_assessment = request.form.get(f'continuous_assessment_{student.id}')
                part_a = request.form.get(f'part_a_{student.id}')
                part_b = request.form.get(f'part_b_{student.id}')
                
                existing_mark.attendance = float(attendance) if attendance else None
                existing_mark.continuous_assessment = float(continuous_assessment) if continuous_assessment else None
                existing_mark.part_a = float(part_a) if part_a else None
                existing_mark.part_b = float(part_b) if part_b else None
                
                total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.continuous_assessment, existing_mark.part_a, existing_mark.part_b]))
            
            elif subject.subject_type == 'Sessional':
                attendance = request.form.get(f'attendance_{student.id}')
                sessional_report = request.form.get(f'sessional_report_{student.id}')
                sessional_viva = request.form.get(f'sessional_viva_{student.id}')

                existing_mark.attendance = float(attendance) if attendance else None
                existing_mark.sessional_report = float(sessional_report) if sessional_report else None
                existing_mark.sessional_viva = float(sessional_viva) if sessional_viva else None

                total_marks = sum(filter(None, [existing_mark.attendance, existing_mark.sessional_report, existing_mark.sessional_viva]))

            elif subject.subject_type == 'Dissertation':
                supervisor_assessment = request.form.get(f'supervisor_assessment_{student.id}')
                proposal_presentation = request.form.get(f'proposal_presentation_{student.id}')
                project_report = request.form.get(f'project_report_{student.id}')
                defense = request.form.get(f'defense_{student.id}')
                
                existing_mark.supervisor_assessment = float(supervisor_assessment) if supervisor_assessment else None
                existing_mark.proposal_presentation = float(proposal_presentation) if proposal_presentation else None
                existing_mark.project_report = float(project_report) if project_report else None
                existing_mark.defense = float(defense) if defense else None

                total_marks = sum(filter(None, [existing_mark.supervisor_assessment, existing_mark.proposal_presentation, existing_mark.project_report, existing_mark.defense]))
            
            existing_mark.total_marks = total_marks
            existing_mark.grade_point, existing_mark.grade_letter = calculate_grade(total_marks, is_retake=is_retake)

        db.session.commit()
        flash(f'Marks for {subject.name} saved successfully!', 'success')
        return redirect(url_for('result_management.add_marks', session_id=session_id, subject_id=subject.id))

    return render_template('rm_add_marks.html', 
                           session=session, 
                           subjects=subjects,
                           students=students,
                           selected_subject=selected_subject,
                           marks_data=marks_data,
                           registrations_data=registrations_data)

@result_management_bp.route('/view_results/<int:session_id>')
@login_required
def view_results(session_id):
    session = RSession.query.get_or_404(session_id)
    return render_template('rm_view_results.html', session=session)

@result_management_bp.route('/course_wise_result/<int:session_id>')
@login_required
def course_wise_result(session_id):
    session = RSession.query.get_or_404(session_id)
    subjects = RSubject.query.filter_by(session_id=session_id).order_by(RSubject.code).all()
    selected_subject_id = request.args.get('subject_id', type=int)
    results = []
    selected_subject = None
    if selected_subject_id:
        selected_subject = RSubject.query.get(selected_subject_id)
        # Define columns to select
        base_columns = [
            RStudent.student_id, RStudent.name, RMark.total_marks,
            RMark.grade_letter, RMark.grade_point, RMark.is_retake
        ]
        extra_columns = []
        if selected_subject:
            if selected_subject.subject_type in ['Theory', 'Theory (UG)']:
                extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
            elif selected_subject.subject_type == 'Sessional':
                extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
            elif selected_subject.subject_type == 'Dissertation':
                if selected_subject.dissertation_type == 'Type1':
                    extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
                else:  # Type2
                    extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
        all_columns = base_columns + extra_columns
        results = db.session.query(*all_columns)\
            .join(RMark, RStudent.id == RMark.student_id)\
            .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == selected_subject_id))\
            .filter(RMark.subject_id == selected_subject_id)\
            .order_by(RStudent.student_id).all()
    return render_template('rm_course_wise_result.html',
                           session=session,
                           subjects=subjects,
                           selected_subject_id=selected_subject_id,
                           selected_subject=selected_subject,
                           results=results)

@result_management_bp.route('/student_wise_result/<int:session_id>')
@login_required
def student_wise_result(session_id):
    session = RSession.query.get_or_404(session_id)
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()

    selected_student_id = request.args.get('student_id', type=int)
    selected_student = None
    results = []
    term_assessment = {}

    if selected_student_id:
        selected_student = RStudent.query.get(selected_student_id)
        # Fetch results for the selected student (ONLY registered courses)
        results = db.session.query(
            RSubject.code.label('subject_code'),
            RSubject.name.label('subject_name'),
            RSubject.credit.label('registered_credits'),
            RMark.grade_letter,
            RMark.grade_point,
            RMark.is_retake,
            RSubject.subject_type
        ).select_from(RMark)\
         .join(RStudent, RStudent.id == RMark.student_id)\
         .join(RSubject, RSubject.id == RMark.subject_id)\
         .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))\
         .filter(RStudent.id == selected_student_id)\
         .order_by(RSubject.code).all()

        total_registered_credits = 0
        total_earned_credits = 0
        total_earned_credit_points = 0
        
        processed_results = []
        for res in results:
            earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
            earned_credit_points = (res.grade_point or 0) * res.registered_credits
            
            processed_results.append({
                'subject_code': res.subject_code,
                'subject_name': res.subject_name,
                'registered_credits': res.registered_credits,
                'grade_letter': res.grade_letter,
                'grade_point': res.grade_point,
                'earned_credits': earned_credits,
                'earned_credit_points': earned_credit_points,
                'remarks': 'Retake' if res.is_retake else ''
            })

            total_registered_credits += res.registered_credits
            total_earned_credits += earned_credits
            total_earned_credit_points += earned_credit_points

        tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
        
        term_assessment = {
            'total_registered_credits': total_registered_credits,
            'total_earned_credits': total_earned_credits,
            'total_earned_credit_points': total_earned_credit_points,
            'tgpa': tgpa
        }
        results = processed_results
        
    return render_template('rm_student_wise_result.html',
                           session=session,
                           students=students,
                           selected_student_id=selected_student_id,
                           selected_student=selected_student,
                           results=results,
                           term_assessment=term_assessment)


@result_management_bp.route('/course_registration/<int:session_id>', methods=['GET', 'POST'])
@login_required
def course_registration(session_id):
    session = RSession.query.get_or_404(session_id)
    subjects = RSubject.query.filter_by(session_id=session_id).order_by(RSubject.code).all()
    students = RStudent.query.filter_by(session_id=session_id).order_by(RStudent.student_id).all()

    selected_subject_id = request.args.get('subject_id', type=int)
    selected_subject = RSubject.query.get(selected_subject_id) if selected_subject_id else None

    registrations = {}
    if selected_subject:
        # Load existing registrations for the selected subject
        existing_regs = db.session.query(RCourseRegistration).filter_by(subject_id=selected_subject.id).all()
        for reg in existing_regs:
            registrations[reg.student_id] = reg

    if request.method == 'POST':
        subject_id = int(request.form.get('subject_id'))
        if not subject_id:
            flash('A subject must be selected.', 'danger')
            return redirect(url_for('result_management.course_registration', session_id=session_id))
        
        # Get all student IDs that were submitted (i.e., whose checkboxes could have been checked)
        students_on_page = RStudent.query.filter_by(session_id=session_id).all()
        student_ids_on_page = {s.id for s in students_on_page}

        # First, delete all existing registrations for this subject for the students shown on the page
        db.session.query(RCourseRegistration).filter(
            RCourseRegistration.subject_id == subject_id,
            RCourseRegistration.student_id.in_(student_ids_on_page)
        ).delete(synchronize_session=False)

        # Now, add back the ones that were checked in the form
        for student_id in student_ids_on_page:
            if f'reg_{student_id}' in request.form:
                is_retake = f'retake_{student_id}' in request.form
                new_reg = RCourseRegistration(
                    student_id=student_id,
                    subject_id=subject_id,
                    is_retake=is_retake
                )
                db.session.add(new_reg)

                # Sync the is_retake flag with the corresponding RMark record
                mark = RMark.query.filter_by(student_id=student_id, subject_id=subject_id).first()
                if mark:
                    mark.is_retake = is_retake
                    # Recalculate grade if total_marks exists
                    if mark.total_marks is not None:
                        mark.grade_point, mark.grade_letter = calculate_grade(mark.total_marks, is_retake=is_retake)
            else:
                 # If a student is unregistered, ensure their mark record also has is_retake as False
                mark = RMark.query.filter_by(student_id=student_id, subject_id=subject_id).first()
                if mark and mark.is_retake:
                    mark.is_retake = False
                    if mark.total_marks is not None:
                        mark.grade_point, mark.grade_letter = calculate_grade(mark.total_marks, is_retake=False)

        db.session.commit()
        flash('Course registration updated successfully!', 'success')
        return redirect(url_for('result_management.course_registration', session_id=session_id, subject_id=subject_id))

    return render_template('rm_course_registration.html', 
                           session=session, 
                           students=students, 
                           subjects=subjects,
                           selected_subject=selected_subject,
                           registrations=registrations)


@result_management_bp.route('/delete_session/<int:session_id>', methods=['POST'])
@login_required
def delete_session(session_id):
    session = RSession.query.get_or_404(session_id)
    db.session.delete(session)
    db.session.commit()
    flash('Session and all related data deleted successfully.', 'success')
    return redirect(url_for('result_management.index'))


class PDFGenerator:
    def __init__(self, buffer, pagesize):
        from reportlab.platypus import SimpleDocTemplate
        from reportlab.lib.units import inch
        self.buffer = buffer
        self.doc = SimpleDocTemplate(
            buffer,
            pagesize=pagesize,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch,
            leftMargin=0.5*inch,
            rightMargin=0.5*inch
        )
        self.story = []

    def _footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 9)
        page_number_text = f"Page {doc.page} of {doc.doc.page_count}"
        canvas.drawRightString(letter[0] - 40, 30, page_number_text)
        canvas.restoreState()

    def build(self, elements):
        # A two-pass approach to get total page numbers for the footer
        
        # First pass
        doc_temp = SimpleDocTemplate(BytesIO(), pagesize=letter)
        frame = Frame(doc_temp.leftMargin, doc_temp.bottomMargin, doc_temp.width, doc_temp.height, id='normal')
        template = PageTemplate(id='main_temp', frames=[frame])
        doc_temp.addPageTemplates([template])
        doc_temp.build(elements)
        self.total_pages = doc_temp.page

        # Second pass (actual build)
        frame = Frame(self.doc.leftMargin, self.doc.bottomMargin, self.doc.width, self.doc.height, id='normal')
        template = PageTemplate(id='main', frames=[frame], onPage=self._footer)
        self.doc.addPageTemplates([template])
        self.doc.build(elements)

class CourseTabulationPDF(PDFGenerator):
    def __init__(self, buffer, subject, session):
        from reportlab.lib.pagesizes import A4
        super().__init__(buffer, pagesize=A4)
        self.subject = subject
        self.session = session
        self.doc.title = f"Course_Result_{subject.code}"

    def generate_elements(self, results):
        styles = getSampleStyleSheet()
        # Custom styles
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='Right', alignment=TA_RIGHT))
        styles.add(ParagraphStyle(name='Left', alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='Line_Data', parent=styles['Normal'], alignment=TA_CENTER, leading=14))
        styles.add(ParagraphStyle(name='TableCellCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_CENTER, leading=8))
        styles.add(ParagraphStyle(name='TableCellLeftCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_LEFT, leading=8))
        styles.add(ParagraphStyle(name='InfoCompact', parent=styles['Normal'], fontSize=9, leading=10))
        
        # Center align headers
        styles['h1'].alignment = TA_CENTER
        styles['h2'].alignment = TA_CENTER
        
        elements = []
        
        # Header
        elements.append(Paragraph("Khulna University", styles['h1']))
        elements.append(Paragraph("Course-wise Tabulation Sheet", styles['h2']))
        elements.append(Spacer(1, 0.08*inch))

        # Info Table (compact)
        info_data = [
            [
                Paragraph(f"<b>Year:</b> {self.session.year or 'N/A'}<br/><b>Discipline:</b> Law<br/><b>Course No.:</b> {self.subject.code}<br/><b>Course Title:</b> {self.subject.name}", styles['InfoCompact']),
                Paragraph(f"<b>Term:</b> {self.session.term}<br/><b>School:</b> Law<br/><b>CH:</b> {self.subject.credit:.1f}<br/><br/><b>Session:</b> {self.session.name}", styles['InfoCompact'])
            ]
        ]
        info_table = Table(info_data, colWidths=[3.6*inch, 2.0*inch])
        info_table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('FONTSIZE', (0,0), (-1,-1), 9)]))
        elements.append(info_table)
        elements.append(Spacer(1, 0.08*inch))

        # Base headers
        base_headers = ['Student\nNo.']
        specific_headers = []
        # Specific headers based on subject type
        if self.subject.subject_type in ['Theory', 'Theory (UG)']:
            specific_headers = ['Attendance\n(10)', 'C.A.\n(40)', 'Sec. A\n(25)', 'Sec. B\n(25)']
        elif self.subject.subject_type == 'Sessional':
            specific_headers = ['Attendance\n(10)', 'Report\n(60)', 'Viva\n(30)']
        elif self.subject.subject_type == 'Dissertation':
            if self.subject.dissertation_type == 'Type1':
                specific_headers = ['Supervisor\n(70)', 'Presentation\n(30)']
            else:
                specific_headers = ['Supervisor\n(50)', 'Report\n(25)', 'Defense\n(25)']
        end_headers = ['Total\nMarks\n(100)', 'Grade\nPoint', 'Grade\nLetter', 'Remarks']
        table_headers = base_headers + specific_headers + end_headers
        data = [[Paragraph(h.replace('\n', '<br/>'), styles['TableCellCompact']) for h in table_headers]]
        for res in results:
            row_data = [Paragraph(str(res.student_id), styles['TableCellCompact'])]
            marks_data = list(res)[6:]
            for mark in marks_data:
                row_data.append(Paragraph(f"{mark:.1f}" if mark is not None else '', styles['TableCellCompact']))
            row_data.extend([
                Paragraph(f"{res.total_marks:.2f}" if res.total_marks is not None else '', styles['TableCellCompact']),
                Paragraph(f"{res.grade_point:.2f}" if res.grade_point is not None else '', styles['TableCellCompact']),
                Paragraph(res.grade_letter or '', styles['TableCellCompact']),
                Paragraph('Retake' if res.is_retake else '', styles['TableCellCompact'])
            ])
            data.append(row_data)
        # Dynamic column widths, reduce by 10%
        base_widths = [1.1*inch] + [0.8*inch]*len(specific_headers) + [0.7*inch, 0.7*inch, 0.7*inch, 0.8*inch]
        col_widths = [w * 0.9 for w in base_widths]
        table = Table(data, colWidths=col_widths, rowHeights=0.35*inch)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 3),
            ('TOPPADDING', (0, 0), (-1, 0), 3),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ])
        table.setStyle(style)
        elements.append(table)
        elements.append(Spacer(1, 0.5*inch))  # Add more space before signature
        
        # Signature section (wider, with space)
        signature_table = Table([
            [Paragraph('-----------------------', styles['Center']), '', Paragraph('-----------------------', styles['Center']), '', Paragraph('-----------------------', styles['Center'])],
            [
                Paragraph('Signature of the First Tabulator<br/>Date:', styles['Center']), '',
                Paragraph('Signature of the Chairman of the Examination Committee<br/>Date:', styles['Center']), '',
                Paragraph('Signature of the Second Tabulator<br/>Date:', styles['Center'])
            ]
        ], colWidths=[2.2*inch, 0.7*inch, 2.2*inch, 0.7*inch, 2.2*inch])
        signature_table.setStyle(TableStyle([
            ('VALIGN', (0, 1), (0, 1), 'CENTER'),
            ('ALIGN', (2, 1), (2, 1), 'CENTER'),
            ('ALIGN', (4, 1), (4, 1), 'CENTER'),
            ('FONTNAME', (0, 1), (4, 1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (4, 1), 8),
            ('TOPPADDING', (0, 1), (4, 1), 10),
        ]))
        elements.append(signature_table)
        
        return elements

class StudentTabulationPDF(PDFGenerator):
    def __init__(self, buffer, student, session):
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate
        pagesize = landscape(A4)
        super().__init__(buffer, pagesize=pagesize)
        self.student = student
        self.session = session
        self.page_count = 1
        self.doc.title = f"Tabulation_{student.student_id}"
        # Restore default margins
        # (no aggressive margin reduction)

    def generate_elements(self, results, term_assessment):
        from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.units import inch
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='Bold', parent=styles['Normal'], fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='TableCell', parent=styles['Normal'], fontSize=10, wordWrap='CJK', alignment=TA_LEFT))
        styles.add(ParagraphStyle(name='TableCellCenter', parent=styles['Normal'], fontSize=10, wordWrap='CJK', alignment=TA_CENTER))
        styles.add(ParagraphStyle(name='TableCellCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_LEFT, leading=8))
        styles.add(ParagraphStyle(name='TableCellCenterCompact', parent=styles['Normal'], fontSize=7, wordWrap='CJK', alignment=TA_CENTER, leading=8))
        styles.add(ParagraphStyle(name='InfoCompact', parent=styles['Normal'], fontSize=9, leading=10))
        elements = []

        # Title
        elements.append(Paragraph("Khulna University", styles['Title']))
        elements.append(Paragraph("Student-wise Tabulation Sheet", styles['Title']))
        elements.append(Spacer(1, 0.08*inch))

        # Two-column info table (compact)
        info_data = [
            [Paragraph('<b>Year:</b>', styles['InfoCompact']), str(self.session.year or self.student.year or 'N/A'),
             Paragraph('<b>Term:</b>', styles['InfoCompact']), str(self.session.term)],
            [Paragraph('<b>Student No.:</b>', styles['InfoCompact']), str(self.student.student_id),
             Paragraph('<b>Name of Student:</b>', styles['InfoCompact']), str(self.student.name)],
            [Paragraph('<b>Discipline:</b>', styles['InfoCompact']), str(self.student.discipline or 'Law'),
             Paragraph('<b>Session:</b>', styles['InfoCompact']), str(self.session.name)],
            [Paragraph('<b>School:</b>', styles['InfoCompact']), 'Law', '', '']
        ]
        info_table = Table(info_data, colWidths=[0.9*inch, 1.7*inch, 0.9*inch, 1.7*inch])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 1),
            ('TOPPADDING', (0,0), (-1,-1), 1),
        ]))
        elements.append(info_table)
        elements.append(Spacer(1, 0.08*inch))

        # If no results, show message and return
        if not results:
            elements.append(Paragraph("No results found for this student.", styles['Normal']))
            return elements

        # Results Table (compact)
        headers = ['Course No.', 'Course Title', 'Registered Credit Hours', 'Letter Grade', 'Grade Point (GP)', 'Earned Credit Hours (CH)', 'Earned Credit Points (GP*CH)', 'Remarks']
        data = [[Paragraph(h, styles['TableCellCenterCompact']) for h in headers]]
        total_registered_credits = 0
        total_earned_credits = 0
        total_earned_credit_points = 0
        for res in results:
            # Always show 0 for GP, Earned CH, Earned Credit Points if missing/None
            grade_point = res.get('grade_point', 0)
            earned_credits = res.get('earned_credits', 0)
            earned_credit_points = res.get('earned_credit_points', 0)
            row = [
                Paragraph(str(res.get('subject_code', '') or ''), styles['TableCellCompact']),
                Paragraph(str(res.get('subject_name', '') or ''), styles['TableCellCompact']),
                Paragraph(str(res.get('registered_credits', '') or ''), styles['TableCellCenterCompact']),
                Paragraph(str(res.get('grade_letter', '') or ''), styles['TableCellCenterCompact']),
                Paragraph(f"{grade_point}" if grade_point is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(f"{earned_credits}" if earned_credits is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(f"{earned_credit_points}" if earned_credit_points is not None else "0", styles['TableCellCenterCompact']),
                Paragraph(str(res.get('remarks', '') or ''), styles['TableCellCenterCompact'])
            ]
            # If mark distribution fields exist, show them (even for F)
            extra_fields = []
            for key in ['attendance', 'continuous_assessment', 'part_a', 'part_b', 'sessional_report', 'sessional_viva', 'supervisor_assessment', 'proposal_presentation', 'project_report', 'defense']:
                if key in res:
                    val = res.get(key, None)
                    extra_fields.append(Paragraph(str(val) if val is not None else '', styles['TableCellCenterCompact']))
            if extra_fields:
                row = row[:3] + extra_fields + row[3:]
            data.append(row)
            total_registered_credits += res['registered_credits'] if res['registered_credits'] else 0
            total_earned_credits += res['earned_credits'] if res['earned_credits'] else 0
            total_earned_credit_points += res['earned_credit_points'] if res['earned_credit_points'] else 0
        # Total row
        data.append([
            '', '',
            Paragraph(f"Total = {total_registered_credits}", styles['TableCellCenterCompact']), '', '',
            Paragraph(f"{total_earned_credits}", styles['TableCellCenterCompact']),
            Paragraph(f"{total_earned_credit_points}", styles['TableCellCenterCompact']),
            ''
        ])
        # Reduce column widths by 20% for compactness
        base_widths = [1.5*inch, 2.5*inch, 1.1*inch, 1.1*inch, 1.2*inch, 1.3*inch, 1.5*inch, 1.1*inch]
        col_widths = [w * 0.8 for w in base_widths]
        table = Table(data, colWidths=col_widths)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 0.7, colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('ALIGN', (1, 1), (1, -1), 'LEFT'),
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ])
        table.setStyle(style)
        elements.append(table)
        elements.append(Spacer(1, 0.08*inch))

        # Term Assessment (compact)
        elements.append(Paragraph('<b>Term Assessment</b>', styles['InfoCompact']))
        elements.append(Paragraph(f"Total Earned Credit Hours in this Term (TCH) = {term_assessment['total_earned_credits']}", styles['InfoCompact']))
        elements.append(Paragraph(f"Total Registered Credit Hours in this Term (RCH) = {term_assessment['total_registered_credits']}", styles['InfoCompact']))
        elements.append(Paragraph(f"Total Earned Credit Points in this Term (TCP) = {term_assessment['total_earned_credit_points']}", styles['InfoCompact']))
        elements.append(Paragraph(f"TGPA = TCP/RCH = {term_assessment['tgpa']:.2f}", styles['InfoCompact']))
        elements.append(Spacer(1, 0.18*inch))

        # Signature lines (compact, with spacing)
        sig_col_width = 1.7*inch
        spacer_col_width = 0.4*inch
        sig_table = Table([
            ['', '', '', '', ''],
            [
                Paragraph('Signature of the First Tabulator<br/>Date:', styles['TableCellCenterCompact']), '',
                Paragraph('Signature of the Chairman of the Examination Committee<br/>Date:', styles['TableCellCenterCompact']), '',
                Paragraph('Signature of the Second Tabulator<br/>Date:', styles['TableCellCenterCompact'])
            ]
        ], colWidths=[sig_col_width, spacer_col_width, sig_col_width, spacer_col_width, sig_col_width])
        sig_table.setStyle(TableStyle([
            ('ALIGN', (0, 1), (0, 1), 'CENTER'),
            ('ALIGN', (2, 1), (2, 1), 'CENTER'),
            ('ALIGN', (4, 1), (4, 1), 'CENTER'),
            ('FONTNAME', (0, 1), (4, 1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (4, 1), 8),
            ('TOPPADDING', (0, 1), (4, 1), 10),
        ]))
        elements.append(sig_table)
        return elements


@result_management_bp.route('/download/course_result/<int:session_id>/<int:subject_id>')
@login_required
def download_course_result(session_id, subject_id):
    subject = RSubject.query.get_or_404(subject_id)
    session = RSession.query.get_or_404(session_id)
    # Define columns to select
    base_columns = [
        RStudent.student_id, RStudent.name, RMark.total_marks,
        RMark.grade_letter, RMark.grade_point, RMark.is_retake
    ]
    extra_columns = []
    if subject:
        if subject.subject_type in ['Theory', 'Theory (UG)']:
            extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
        elif subject.subject_type == 'Sessional':
            extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
        elif subject.subject_type == 'Dissertation':
            if subject.dissertation_type == 'Type1':
                extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
            else:  # Type2
                extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
    all_columns = base_columns + extra_columns
    results = db.session.query(*all_columns)\
        .join(RMark, RStudent.id == RMark.student_id)\
        .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == subject_id))\
        .filter(RMark.subject_id == subject_id)\
        .order_by(RStudent.student_id).all()
    buffer = BytesIO()
    pdf = CourseTabulationPDF(buffer, subject, session)
    elements = pdf.generate_elements(results)
    pdf.doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'Course_{subject.code}_Result.pdf', mimetype='application/pdf')

@result_management_bp.route('/download/student_result/<int:session_id>/<int:student_id>')
@login_required
def download_student_result(session_id, student_id):
    student = RStudent.query.get_or_404(student_id)
    session = RSession.query.get_or_404(session_id)

    # Only include subjects where the student is registered
    results_query = db.session.query(
        RSubject.code.label('subject_code'),
        RSubject.name.label('subject_name'),
        RSubject.credit.label('registered_credits'),
        RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type,
        RMark.total_marks,
        RMark.grade_point, RMark.grade_letter,
        RMark.is_retake,
        RMark.total_marks,
        RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b,
        RMark.sessional_report, RMark.sessional_viva,
        RMark.supervisor_assessment, RMark.proposal_presentation, RMark.project_report, RMark.defense
    ).select_from(RMark)
    results_query = results_query.join(RStudent, RStudent.id == RMark.student_id)
    results_query = results_query.join(RSubject, RSubject.id == RMark.subject_id)
    results_query = results_query.join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))
    results_query = results_query.filter(RStudent.id == student_id)
    results_query = results_query.order_by(RSubject.code)
    results = results_query.all()

    total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
    processed_results = []
    for res in results:
        earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
        earned_credit_points = (res.grade_point or 0) * res.registered_credits
        processed_results.append({
            'subject_code': res.subject_code, 'subject_name': res.subject_name,
            'registered_credits': res.registered_credits, 'grade_letter': res.grade_letter,
            'grade_point': res.grade_point, 'earned_credits': earned_credits,
            'earned_credit_points': earned_credit_points,
            'remarks': 'Retake' if res.is_retake else ''
        })
        total_registered_credits += res.registered_credits
        total_earned_credits += earned_credits
        total_earned_credit_points += earned_credit_points

    tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
    term_assessment = {
        'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
        'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
    }

    buffer = BytesIO()
    pdf = StudentTabulationPDF(buffer, student, session)
    elements = pdf.generate_elements(processed_results, term_assessment)
    pdf.doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'Student_{student.student_id}_Tabulation.pdf', mimetype='application/pdf')

@result_management_bp.route('/download/student_result_docx/<int:session_id>/<int:student_id>')
@login_required
def download_student_result_docx(session_id, student_id):
    student = RStudent.query.get_or_404(student_id)
    session = RSession.query.get_or_404(session_id)

    # Only include subjects where the student is registered
    results_query = db.session.query(
        RSubject.code.label('subject_code'),
        RSubject.name.label('subject_name'),
        RSubject.credit.label('registered_credits'),
        RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type
    ).select_from(RMark)\
     .join(RStudent, RStudent.id == RMark.student_id)\
     .join(RSubject, RSubject.id == RMark.subject_id)\
     .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))\
     .filter(RStudent.id == student_id)\
     .order_by(RSubject.code).all()

    total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
    processed_results = []
    for res in results_query:
        earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
        earned_credit_points = (res.grade_point or 0) * res.registered_credits
        processed_results.append({
            'subject_code': res.subject_code, 'subject_name': res.subject_name,
            'registered_credits': res.registered_credits, 'grade_letter': res.grade_letter,
            'grade_point': res.grade_point, 'earned_credits': earned_credits,
            'earned_credit_points': earned_credit_points,
            'remarks': 'Retake' if res.is_retake else ''
        })
        total_registered_credits += res.registered_credits
        total_earned_credits += earned_credits
        total_earned_credit_points += earned_credit_points

    tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
    term_assessment = {
        'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
        'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
    }

    # Create DOCX
    doc = Document()
    # Title
    doc.add_heading('Khulna University', 0).alignment = 1
    doc.add_heading('Student-wise Tabulation Sheet', level=1).alignment = 1

    # Two-column info table
    info_table = doc.add_table(rows=2, cols=4)
    info_table.autofit = False
    info_table.columns[0].width = Inches(1.2)
    info_table.columns[1].width = Inches(2.2)
    info_table.columns[2].width = Inches(1.2)
    info_table.columns[3].width = Inches(2.2)
    # Left column
    info_table.cell(0,0).text = 'Year:'
    info_table.cell(0,1).text = str(session.year or student.year or 'N/A')
    info_table.cell(1,0).text = 'Student No.:'
    info_table.cell(1,1).text = str(student.student_id)
    # Right column
    info_table.cell(0,2).text = 'Term:'
    info_table.cell(0,3).text = str(session.term)
    info_table.cell(1,2).text = 'Name of Student:'
    info_table.cell(1,3).text = str(student.name)
    # Next row for discipline/session/school
    row = info_table.add_row().cells
    row[0].text = 'Discipline:'
    row[1].text = str(student.discipline or 'Law')
    row[2].text = 'Session:'
    row[3].text = str(session.name)
    row2 = info_table.add_row().cells
    row2[0].text = 'School:'
    row2[1].text = 'Law'
    row2[2].text = ''
    row2[3].text = ''
    for row in info_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.runs[0].font.size = Pt(11)
    doc.add_paragraph('')

    # Results Table
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    headers = ['Course No.', 'Course Title', 'Registered Credit Hours', 'Letter Grade', 'Grade Point (GP)', 'Earned Credit Hours (CH)', 'Earned Credit Points (GP*CH)', 'Remarks']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.runs[0].font.bold = True
            p.alignment = 1
    for result in processed_results:
        row_cells = table.add_row().cells
        row_cells[0].text = str(result['subject_code'])
        row_cells[1].text = str(result['subject_name'])
        row_cells[2].text = f"{result['registered_credits']}"
        row_cells[3].text = result['grade_letter'] or ''
        row_cells[4].text = f"{result['grade_point']}" if result['grade_point'] is not None else ''
        row_cells[5].text = f"{result['earned_credits']}"
        row_cells[6].text = f"{result['earned_credit_points']}"
        row_cells[7].text = result['remarks']
    # Total row
    total_row = table.add_row().cells
    total_row[0].text = ''
    total_row[1].text = ''
    total_row[2].text = f"Total = {total_registered_credits}"
    total_row[3].text = ''
    total_row[4].text = ''
    total_row[5].text = f"{total_earned_credits}"
    total_row[6].text = f"{total_earned_credit_points}"
    total_row[7].text = ''
    for cell in total_row:
        for p in cell.paragraphs:
            p.runs[0].font.bold = True
    doc.add_paragraph('')

    # Term Assessment
    p = doc.add_paragraph()
    p.add_run('Term Assessment').bold = True
    doc.add_paragraph(f"Total Earned Credit Hours in this Term (TCH) = {total_earned_credits}")
    doc.add_paragraph(f"Total Registered Credit Hours in this Term (RCH) = {total_registered_credits}")
    doc.add_paragraph(f"Total Earned Credit Points in this Term (TCP) = {total_earned_credit_points}")
    doc.add_paragraph(f"TGPA = TCP/RCH = {tgpa:.2f}")
    doc.add_paragraph('')

    # Signature lines (3 columns)
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.autofit = False
    sig_table.columns[0].width = Inches(2.5)
    sig_table.columns[1].width = Inches(2.5)
    sig_table.columns[2].width = Inches(2.5)
    sig_table.cell(0,0).text = ''
    sig_table.cell(0,1).text = ''
    sig_table.cell(0,2).text = ''
    sig_table.cell(1,0).text = 'Signature of the First Tabulator\nDate:'
    sig_table.cell(1,1).text = 'Signature of the Second Tabulator\nDate:'
    sig_table.cell(1,2).text = 'Signature of the Chairman, Examination Committee\nDate:'
    for i in range(3):
        for p in sig_table.cell(1,i).paragraphs:
            p.alignment = 1
    doc.add_paragraph('')

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f'Student_{student.student_id}_Tabulation.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@result_management_bp.route('/download/all_student_results/<int:session_id>')
@login_required
def download_all_student_results(session_id):
    session = RSession.query.get_or_404(session_id)
    students = RStudent.query.filter_by(session_id=session_id).all()
    
    if not students:
        flash('No students in this session to generate results for.', 'warning')
        return redirect(url_for('result_management.student_wise_result', session_id=session_id))

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zf:
        for student in students:
            # Re-use logic from download_student_result
            results_query = db.session.query(
                RSubject.code.label('subject_code'), RSubject.name.label('subject_name'), RSubject.credit.label('registered_credits'),
                RMark.grade_letter, RMark.grade_point, RMark.is_retake, RSubject.subject_type
            ).select_from(RMark)
            results_query = results_query.join(RStudent, RStudent.id == RMark.student_id)
            results_query = results_query.join(RSubject, RSubject.id == RMark.subject_id)
            results_query = results_query.join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == RSubject.id))
            results_query = results_query.filter(RStudent.id == student.id).order_by(RSubject.code)
            results_query = results_query.all()

            if not results_query: continue

            total_registered_credits, total_earned_credits, total_earned_credit_points = 0, 0, 0
            processed_results = []
            for res in results_query:
                earned_credits = res.registered_credits if (res.grade_point or 0) >= 2.0 else 0
                earned_credit_points = (res.grade_point or 0) * res.registered_credits
                processed_results.append({
                    'subject_code': res.subject_code, 'subject_name': res.subject_name, 'registered_credits': res.registered_credits,
                    'grade_letter': res.grade_letter, 'grade_point': res.grade_point, 'earned_credits': earned_credits,
                    'earned_credit_points': earned_credit_points, 'remarks': 'Retake' if res.is_retake else ''
                })
                total_registered_credits += res.registered_credits
                total_earned_credits += earned_credits
                total_earned_credit_points += earned_credit_points
            
            tgpa = total_earned_credit_points / total_registered_credits if total_registered_credits > 0 else 0
            term_assessment = {
                'total_registered_credits': total_registered_credits, 'total_earned_credits': total_earned_credits,
                'total_earned_credit_points': total_earned_credit_points, 'tgpa': tgpa
            }
            
            pdf_buffer = BytesIO()
            pdf = StudentTabulationPDF(pdf_buffer, student, session)
            elements = pdf.generate_elements(processed_results, term_assessment)
            pdf.doc.build(elements)
            pdf_buffer.seek(0)
            zf.writestr(f'Student_{student.student_id}_Tabulation.pdf', pdf_buffer.read())

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'All_Student_Results_{session.name}.zip', mimetype='application/zip')

@result_management_bp.route('/download/all_course_results/<int:session_id>')
@login_required
def download_all_course_results(session_id):
    session = RSession.query.get_or_404(session_id)
    subjects = RSubject.query.filter_by(session_id=session_id).all()

    if not subjects:
        flash('No subjects in this session to generate results for.', 'warning')
        return redirect(url_for('result_management.course_wise_result', session_id=session_id))
        
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zf:
        for subject in subjects:
            base_columns = [
                RStudent.student_id, RStudent.name, RMark.total_marks,
                RMark.grade_letter, RMark.grade_point, RMark.is_retake
            ]
            extra_columns = []
            if subject.subject_type in ['Theory', 'Theory (UG)']:
                extra_columns = [RMark.attendance, RMark.continuous_assessment, RMark.part_a, RMark.part_b]
            elif subject.subject_type == 'Sessional':
                extra_columns = [RMark.attendance, RMark.sessional_report, RMark.sessional_viva]
            elif subject.subject_type == 'Dissertation':
                if subject.dissertation_type == 'Type1':
                    extra_columns = [RMark.supervisor_assessment, RMark.proposal_presentation]
                else:
                    extra_columns = [RMark.supervisor_assessment, RMark.project_report, RMark.defense]
            all_columns = base_columns + extra_columns
            results = db.session.query(*all_columns)\
                .join(RMark, RStudent.id == RMark.student_id)\
                .join(RCourseRegistration, (RCourseRegistration.student_id == RStudent.id) & (RCourseRegistration.subject_id == subject.id))\
                .filter(RMark.subject_id == subject.id)\
                .order_by(RStudent.student_id).all()

            if not results: continue

            pdf_buffer = BytesIO()
            pdf = CourseTabulationPDF(pdf_buffer, subject, session)
            elements = pdf.generate_elements(results)
            pdf.doc.build(elements)
            pdf_buffer.seek(0)
            zf.writestr(f'Course_{subject.code}_Result.pdf', pdf_buffer.read())

    zip_buffer.seek(0)
    return send_file(zip_buffer, as_attachment=True, download_name=f'All_Course_Results_{session.name}.zip', mimetype='application/zip')
